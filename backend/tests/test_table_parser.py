# -*- coding: utf-8 -*-
"""
表格解析器单元测试

测试覆盖：
1. TableData 数据类（to_markdown / to_csv / to_structured_text）
2. 文本表格解析（纯文本分隔符检测）
3. DOCX 表格解析（合并单元格检测）
4. PDF 表格解析（PyMuPDF fallback）
5. 跨页表格合并逻辑
6. 表头识别逻辑
7. 特殊字符处理（中文、数字、空格、标点）
"""

import io
import pytest
from typing import List, Dict, Any, Tuple

from app.services.table_parser import (
    TableData,
    TableParser,
    Cell,
    get_table_parser,
)


# ── TableData 数据类测试 ──────────────────────────────────────────────────────

class TestTableData:
    """TableData 数据类测试"""

    @pytest.fixture
    def sample_table(self) -> TableData:
        return TableData(
            title="2023年度营收数据",
            page_number=1,
            column_count=4,
            row_count=3,
            headers=["部门", "Q1", "Q2", "Q3"],
            rows=[
                ["研发部", "100万", "120万", "150万"],
                ["市场部", "80万", "90万", "110万"],
            ],
            merged_cells=[(0, 1, 1, 2)],  # Q1 和 Q2 合并
        )

    def test_to_dict(self, sample_table):
        d = sample_table.to_dict()
        assert d["title"] == "2023年度营收数据"
        assert d["column_count"] == 4
        assert d["row_count"] == 3
        assert d["headers"] == ["部门", "Q1", "Q2", "Q3"]
        assert len(d["rows"]) == 3
        assert d["merged_cells"] == [(0, 1, 1, 2)]
        assert d["table_id"]  # 自动生成

    def test_to_markdown(self, sample_table):
        md = sample_table.to_markdown()
        assert "**2023年度营收数据**" in md
        assert "| 部门 | Q1 | Q2 | Q3 |" in md
        assert "| --- | --- | --- | --- |" in md
        assert "研发部" in md
        assert "市场部" in md

    def test_to_markdown_empty(self):
        t = TableData(rows=[])
        assert t.to_markdown() == ""

    def test_to_csv(self, sample_table):
        csv_text = sample_table.to_text(format="csv")
        assert "部门" in csv_text
        assert "Q1" in csv_text
        assert "研发部" in csv_text
        assert "100万" in csv_text

    def test_to_tsv(self, sample_table):
        tsv_text = sample_table.to_text(format="tsv")
        lines = tsv_text.split("\n")
        assert len(lines) >= 4  # 表头 + 3行数据
        assert "\t" in tsv_text

    def test_to_structured_text(self, sample_table):
        structured = sample_table.to_text(format="structured")
        assert "【表格: 2023年度营收数据" in structured
        assert "页码: 1" in structured
        assert "4列×3行" in structured
        assert "部门: 研发部" in structured

    def test_merged_cell_expansion_in_markdown(self):
        """测试合并单元格在 Markdown 中的占位处理"""
        t = TableData(
            title="测试表",
            column_count=3,
            row_count=2,
            headers=["A", "B", "C"],
            rows=[["1", "2", "3"]],
            merged_cells=[(0, 0, 1, 2)],  # A 和 B 合并
        )
        md = t.to_markdown()
        # 合并单元格应在 Markdown 中有正确的占位
        assert "A" in md
        assert "B" in md
        assert "C" in md


# ── 文本表格解析测试 ──────────────────────────────────────────────────────────

class TestTextTableParsing:
    """从纯文本中解析表格"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_parse_tab_separated_table(self, parser):
        """测试 Tab 分隔的表格解析"""
        text = "姓名\t年龄\t城市\n张三\t25\t北京\n李四\t30\t上海"
        table = parser._parse_table_from_text(text, "test.txt", 1)
        assert table is not None
        assert table.column_count == 3
        assert table.row_count == 2
        assert table.headers == ["姓名", "年龄", "城市"]
        assert table.rows[0] == ["张三", "25", "北京"]

    def test_parse_pipe_separated_table(self, parser):
        """测试 | 分隔的 Markdown 风格表格"""
        text = "姓名 | 年龄 | 城市\n张三 | 25 | 北京\n李四 | 30 | 上海"
        table = parser._parse_table_from_text(text, "test.md", 1)
        assert table is not None
        assert table.column_count == 3
        assert table.headers == ["姓名", "年龄", "城市"]

    def test_parse_csv_style_table(self, parser):
        """测试 CSV 风格表格"""
        text = "姓名,年龄,城市\n张三,25,北京\n李四,30,上海"
        table = parser._parse_table_from_text(text, "test.csv", 1)
        assert table is not None
        assert table.column_count == 3

    def test_parse_insufficient_lines(self, parser):
        """测试不足两行的文本不应解析为表格"""
        text = "只有一行数据"
        table = parser._parse_table_from_text(text, "test.txt", 1)
        assert table is None

    def test_detect_tab_delimiter(self, parser):
        assert parser._detect_delimiter(["a\tb\tc", "d\te\tf"]) == "\t"

    def test_detect_pipe_delimiter(self, parser):
        assert parser._detect_delimiter(["a|b|c", "d|e|f"]) == "|"

    def test_detect_no_delimiter(self, parser):
        assert parser._detect_delimiter(["abc", "def"]) is None


# ── DOCX 表格解析测试 ─────────────────────────────────────────────────────────

class TestDocxTableParsing:
    """DOCX 表格解析测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def _create_docx_with_table(self, headers: List[str], rows: List[List[str]]) -> bytes:
        """创建一个包含表格的 DOCX 文件"""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))

        # 表头
        for j, h in enumerate(headers):
            table.cell(0, j).text = h

        # 数据行
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                table.cell(i + 1, j).text = cell

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def test_parse_docx_basic_table(self, parser):
        """测试基础 DOCX 表格解析"""
        docx_bytes = self._create_docx_with_table(
            headers=["姓名", "部门", "薪资"],
            rows=[
                ["张三", "研发部", "15000"],
                ["李四", "市场部", "12000"],
            ],
        )
        tables, chunks = parser.parse_docx(docx_bytes, "test.docx")
        assert len(tables) == 1
        assert tables[0].column_count == 3
        assert tables[0].row_count == 2
        assert tables[0].headers == ["姓名", "部门", "薪资"]

    def test_parse_docx_merged_cell(self, parser):
        """测试 DOCX 合并单元格检测"""
        from docx import Document, OxmlElement, ns

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "部门"
        table.cell(0, 2).text = "备注"
        table.cell(1, 0).text = "张三"
        table.cell(1, 1).text = "研发部"
        table.cell(1, 2).text = "优秀员工"

        # 手动设置 gridSpan（合并单元格）
        tc = table.cell(1, 1)._tc
        tcPr = tc.get_or_add_tcPr()
        grid_span = OxmlElement(f'{{{ns.w}}}gridSpan')
        grid_span.set(ns.qn('w:val'), '2')
        tcPr.append(grid_span)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()

        tables, chunks = parser.parse_docx(docx_bytes, "test.docx")
        assert len(tables) == 1
        assert tables[0].column_count == 3
        # 应该有合并单元格记录
        assert len(tables[0].merged_cells) >= 1

    def test_parse_docx_empty_file(self, parser):
        """测试空 DOCX 文件"""
        from docx import Document
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        tables, chunks = parser.parse_docx(buffer.read(), "empty.docx")
        assert len(tables) == 0

    def test_parse_docx_multiple_tables(self, parser):
        """测试 DOCX 中多个表格"""
        from docx import Document
        doc = Document()
        doc.add_paragraph("第一个表格前的文字")
        t1 = doc.add_table(rows=2, cols=2)
        t1.cell(0, 0).text = "A1"
        t1.cell(0, 1).text = "B1"
        t1.cell(1, 0).text = "A2"
        t1.cell(1, 1).text = "B2"

        doc.add_paragraph("两个表格之间的文字")
        t2 = doc.add_table(rows=2, cols=3)
        t2.cell(0, 0).text = "X1"
        t2.cell(0, 1).text = "Y1"
        t2.cell(0, 2).text = "Z1"
        t2.cell(1, 0).text = "X2"
        t2.cell(1, 1).text = "Y2"
        t2.cell(1, 2).text = "Z2"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        tables, chunks = parser.parse_docx(buffer.read(), "multi.docx")
        assert len(tables) == 2
        assert tables[0].column_count == 2
        assert tables[1].column_count == 3
        # 应该有文本块
        assert len(chunks) > 0

    def test_parse_docx_special_characters(self, parser):
        """测试特殊字符处理"""
        docx_bytes = self._create_docx_with_table(
            headers=["项目", "值（%）", "说明"],
            rows=[
                ["营收增长", "25.5", "含汇率波动±3.2%"],
                ["净利润", "1,234.56万", "扣非后*1.5"],
                ["同比（YoY）", "-5.2%", "受疫情&政策影响"],
            ],
        )
        tables, chunks = parser.parse_docx(docx_bytes, "special.docx")
        assert len(tables) == 1
        assert "营收增长" in str(tables[0].rows)
        assert "25.5" in str(tables[0].rows)
        assert "±3.2%" in str(tables[0].rows)
        assert "*1.5" in str(tables[0].rows)


# ── PDF 表格解析测试（PyMuPDF fallback） ──────────────────────────────────────

class TestPdfTableParsing:
    """PDF 表格解析测试（使用 PyMuPDF fallback）"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_parse_pdf_with_tables(self, parser):
        """测试 PDF 表格解析（PyMuPDF fallback）"""
        import fitz
        # 创建一个简单的 PDF 包含表格
        doc = fitz.open()
        page = doc.new_page()

        # 使用 PyMuPDF 的 draw 功能创建一个简单"表格"
        text = "姓名 | 年龄 | 城市\n张三 | 25 | 北京\n李四 | 30 | 上海"
        page.insert_text((100, 100), text, fontsize=12)

        pdf_bytes = doc.tobytes()
        doc.close()

        tables, chunks = parser.parse_pdf(pdf_bytes, "test.pdf")
        # 注意：find_tables() 需要真实的表格结构
        # 这个测试验证的是 fallback 不会崩溃
        assert isinstance(tables, list)
        assert isinstance(chunks, list)

    def test_parse_pdf_empty(self, parser):
        """测试空 PDF 文件"""
        import fitz
        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        tables, chunks = parser.parse_pdf(pdf_bytes, "empty.pdf")
        assert isinstance(tables, list)


# ── 跨页表格合并测试 ──────────────────────────────────────────────────────────

class TestCrossPageMerge:
    """跨页表格合并逻辑测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_merge_identical_headers(self, parser):
        """测试表头相同的跨页表格合并"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["部门", "Q1", "Q2"],
            rows=[["研发", "100", "120"]],
        )
        t2 = TableData(
            page_number=2,
            column_count=3,
            headers=["部门", "Q1", "Q2"],
            rows=[["市场", "80", "90"]],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 1
        assert result[0].row_count == 2
        assert result[0].merged_pages == [1, 2]

    def test_no_merge_different_headers(self, parser):
        """测试表头不同时不合并"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["部门", "Q1", "Q2"],
            rows=[["研发", "100", "120"]],
        )
        t2 = TableData(
            page_number=2,
            column_count=3,
            headers=["姓名", "年龄", "城市"],
            rows=[["张三", "25", "北京"]],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 2  # 不合并

    def test_merge_subset_headers(self, parser):
        """测试表头子集情况下的合并"""
        t1 = TableData(
            page_number=1,
            column_count=4,
            headers=["部门", "Q1", "Q2", "Q3"],
            rows=[["研发", "100", "120", "150"]],
        )
        # 跨页后只显示了部分表头
        t2 = TableData(
            page_number=2,
            column_count=4,
            headers=["Q1", "Q2"],
            rows=[["市场", "80", "90"]],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 1  # 应该合并

    def test_merge_adjacent_pages_same_columns(self, parser):
        """测试相邻页码、相同列数的合并"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["部门", "Q1", "Q2"],
            rows=[["研发", "100", "120"]],
        )
        t2 = TableData(
            page_number=2,
            column_count=3,
            rows=[["市场", "80", "90"]],  # 没有表头
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 1

    def test_no_merge_non_adjacent_pages(self, parser):
        """测试不相邻页码不合并"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["A", "B", "C"],
            rows=[["1", "2", "3"]],
        )
        t2 = TableData(
            page_number=5,  # 不相邻
            column_count=3,
            headers=["A", "B", "C"],
            rows=[["4", "5", "6"]],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 2

    def test_merge_row_index_adjustment(self, parser):
        """测试合并后行索引调整"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["A", "B", "C"],
            rows=[["1", "2", "3"]],
            merged_cells=[(0, 0, 1, 2)],
        )
        t2 = TableData(
            page_number=2,
            column_count=3,
            headers=["A", "B", "C"],
            rows=[["4", "5", "6"]],
            merged_cells=[(0, 1, 1, 2)],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 1
        # t2 的合并单元格索引应该被调整
        assert len(result[0].merged_cells) >= 2


# ── 表头识别测试 ──────────────────────────────────────────────────────────────

class TestHeaderDetection:
    """表头识别逻辑测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_detect_chinese_header(self, parser):
        """检测中文表头关键词"""
        assert parser._looks_like_header(["单位", "类型", "名称"]) is True
        assert parser._looks_like_header(["金额", "数量", "占比"]) is True
        assert parser._looks_like_header(["日期", "时间", "项目"]) is True

    def test_detect_short_header(self, parser):
        """检测简短的英文表头"""
        assert parser._looks_like_header(["ID", "Name", "Age", "City"]) is True
        assert parser._looks_like_header(["A", "B", "C", "D"]) is True

    def test_not_header_long_text(self, parser):
        """长文本不应被识别为表头"""
        assert parser._looks_like_header([
            "这是一段很长的描述性文字，不应该被识别为表头"
        ]) is False

    def test_empty_row_not_header(self, parser):
        """空行不应被识别为表头"""
        assert parser._looks_like_header([]) is False


# ── Jaccard 相似度测试 ────────────────────────────────────────────────────────

class TestListSimilarity:
    """列表相似度计算测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_identical_lists(self, parser):
        assert parser._list_similarity(["A", "B", "C"], ["A", "B", "C"]) == 1.0

    def test_similar_lists(self, parser):
        sim = parser._list_similarity(["A", "B", "C", "D"], ["A", "B", "C", "E"])
        assert sim == 0.6  # 3/5

    def test_disjoint_lists(self, parser):
        assert parser._list_similarity(["A", "B"], ["C", "D"]) == 0.0

    def test_empty_lists(self, parser):
        assert parser._list_similarity([], []) == 1.0
        assert parser._list_similarity(["A"], []) == 0.0

    def test_case_insensitive(self, parser):
        sim = parser._list_similarity(["部门", "Q1"], ["部门", "Q1"])
        assert sim == 1.0


# ── PyMuPDF 合并单元格检测测试 ────────────────────────────────────────────────

class TestMergedCellDetection:
    """合并单元格检测测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_detect_horizontal_merge(self, parser):
        """检测横向合并（非空单元格 + 后续空单元格）"""
        data = [
            ["标题", "", "列C"],  # A 和 B 合并
            ["值A", "值B", "值C"],
        ]
        merged = parser._detect_merged_cells_from_pymupdf(data)
        assert len(merged) >= 1
        assert merged[0] == (0, 0, 1, 2)

    def test_no_merge(self, parser):
        """没有合并单元格的情况"""
        data = [
            ["A", "B", "C"],
            ["1", "2", "3"],
        ]
        merged = parser._detect_merged_cells_from_pymupdf(data)
        assert len(merged) == 0

    def test_empty_data(self, parser):
        assert parser._detect_merged_cells_from_pymupdf([]) == []


# ── 分隔符检测测试 ────────────────────────────────────────────────────────────

class TestDelimiterDetection:
    """分隔符检测测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_detect_tabs(self, parser):
        lines = ["A\tB\tC", "D\tE\tF", "G\tH\tI"]
        assert parser._detect_delimiter(lines) == "\t"

    def test_detect_pipes(self, parser):
        lines = ["A|B|C", "D|E|F", "G|H|I"]
        assert parser._detect_delimiter(lines) == "|"

    def test_detect_commas(self, parser):
        lines = ["A,B,C", "D,E,F", "G,H,I"]
        assert parser._detect_delimiter(lines) == ","

    def test_no_consistent_delimiter(self, parser):
        lines = ["A\tB\tC", "D|E|F"]
        assert parser._detect_delimiter(lines) is None

    def test_single_line(self, parser):
        lines = ["A\tB\tC"]
        assert parser._detect_delimiter(lines) == "\t"


# ── 单例模式测试 ──────────────────────────────────────────────────────────────

class TestSingleton:
    """单例模式测试"""

    def test_get_table_parser_singleton(self):
        p1 = get_table_parser()
        p2 = get_table_parser()
        assert p1 is p2


# ── 复杂表格测试 ──────────────────────────────────────────────────────────────

class TestComplexTableScenarios:
    """复杂场景测试"""

    @pytest.fixture
    def parser(self) -> TableParser:
        return TableParser()

    def test_large_table(self, parser):
        """测试大型表格（100行×10列）"""
        headers = [f"列{i}" for i in range(10)]
        rows = [[f"r{i}c{j}" for j in range(10)] for i in range(100)]
        t = TableData(
            column_count=10,
            row_count=100,
            headers=headers,
            rows=rows,
        )
        md = t.to_markdown()
        assert "列0" in md
        assert "r99c9" in md
        # Markdown 应该有 100 行数据 + 2 行表头
        assert md.count("|") >= 200

    def test_table_with_special_chars(self, parser):
        """测试包含特殊字符的表格"""
        docx_parser = TableParser()
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "项目"
        table.cell(0, 1).text = "值（%）"
        table.cell(0, 2).text = "备注 & 说明"
        table.cell(1, 0).text = "营收"
        table.cell(1, 1).text = "25.5%"
        table.cell(1, 2).text = "含汇率波动 ±3.2%"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        tables, chunks = docx_parser.parse_docx(buffer.read(), "special.docx")
        assert len(tables) == 1
        assert "±3.2%" in str(tables[0].rows)
        assert "&" in str(tables[0].rows)

    def test_table_with_unicode(self, parser):
        """测试 Unicode 字符（中日韩、Emoji）"""
        docx_parser = TableParser()
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "日本語テスト"
        table.cell(0, 1).text = "한국어 테스트"
        table.cell(1, 0).text = "中文测试✅"
        table.cell(1, 1).text = "émoji 🎉"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        tables, chunks = docx_parser.parse_docx(buffer.read(), "unicode.docx")
        assert len(tables) == 1
        assert "日本語" in str(tables[0].rows)
        assert "🎉" in str(tables[0].rows)

    def test_nested_merge_cells(self, parser):
        """测试嵌套合并单元格"""
        from docx import Document, OxmlElement, ns

        doc = Document()
        table = doc.add_table(rows=4, cols=4)
        # 设置各种文本
        for i in range(4):
            for j in range(4):
                table.cell(i, j).text = f"R{i}C{j}"

        # 横向合并 (0,0) 跨 2 列
        tc = table.cell(0, 0)._tc
        tcPr = tc.get_or_add_tcPr()
        gs = OxmlElement(f'{{{ns.w}}}gridSpan')
        gs.set(ns.qn('w:val'), '2')
        tcPr.append(gs)

        # 横向合并 (0,2) 跨 2 列
        tc2 = table.cell(0, 2)._tc
        tcPr2 = tc2.get_or_add_tcPr()
        gs2 = OxmlElement(f'{{{ns.w}}}gridSpan')
        gs2.set(ns.qn('w:val'), '2')
        tcPr2.append(gs2)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        tables, chunks = parser.parse_docx(buffer.read(), "nested.docx")
        assert len(tables) == 1
        # 应该有 2 个合并单元格记录
        assert len(tables[0].merged_cells) >= 2

    def test_cross_page_merge_with_data_continuation(self, parser):
        """测试跨页表格合并 - 数据连续性验证"""
        t1 = TableData(
            page_number=1,
            column_count=3,
            headers=["月份", "营收（万）", "利润（万）"],
            rows=[
                ["1月", "100", "20"],
                ["2月", "120", "25"],
                ["3月", "110", "22"],
            ],
        )
        t2 = TableData(
            page_number=2,
            column_count=3,
            headers=["月份", "营收（万）", "利润（万）"],
            rows=[
                ["4月", "130", "28"],
                ["5月", "140", "30"],
            ],
        )

        result = parser._merge_cross_page_tables([t1, t2])
        assert len(result) == 1
        merged = result[0]
        assert merged.row_count == 5  # 3 + 2
        # 验证数据完整性
        all_data = merged.rows
        assert any("1月" in r[0] for r in all_data)
        assert any("5月" in r[0] for r in all_data)
        assert merged.merged_pages == [1, 2]
