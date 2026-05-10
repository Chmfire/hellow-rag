# -*- coding: utf-8 -*-
"""
表格解析服务（基于 Unstructured + PyMuPDF）

核心能力：
- 多格式支持：PDF / DOCX / PPTX / 图片（PNG/JPG）
- 智能表格边界识别
- 合并单元格处理（行/列跨距）
- 跨页表格自动拼接
- 表格标题/元数据提取
- 结构化输出（JSON / 文本）

输出格式：
  TableData:
    - title: 表格标题（可能为空）
    - page_number: 所在页码
    - bbox: 表格在页面中的边界框 (x0, y0, x1, y1)
    - headers: 表头列表
    - rows: 数据行列表（每行是单元格列表）
    - cell_spans: 合并单元格信息 (row_idx, col_idx, row_span, col_span)
    - column_count: 列数
    - row_count: 行数
    - content_text: 文本化后的表格内容（Markdown 格式）
    - metadata: 额外元数据
"""

import io
import re
import logging
import uuid
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# ── 正则 ──────────────────────────────────────────────────────────────────────

_IMAGE_RE = re.compile(r'<<IMAGE:[0-9a-f]+>>')
_TABLE_TITLE_RE = re.compile(
    r'^[第]?[一二三四五六七八九十\d]+[章章节篇部]?[\s：:,，]*(?:表|表格)\s*[\d]*\s*[:：-]\s*(.+)$'
)


# ── 数据类 ────────────────────────────────────────────────────────────────────

@dataclass
class Cell:
    """单个单元格"""
    text: str = ""
    row_span: int = 1
    col_span: int = 1
    is_header: bool = False
    bbox: Optional[Tuple[float, float, float, float]] = None  # (x0, y0, x1, y1)


@dataclass
class TableData:
    """解析后的表格数据"""
    table_id: str = field(default_factory=lambda: str(uuid.uuid4())[:12])
    title: str = ""
    page_number: int = 0
    file_name: str = ""
    bbox: Optional[Tuple[float, float, float, float]] = None
    column_count: int = 0
    row_count: int = 0
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    # 合并单元格信息：(row_idx, col_idx, row_span, col_span)
    merged_cells: List[Tuple[int, int, int, int]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "table_id": self.table_id,
            "title": self.title,
            "page_number": self.page_number,
            "bbox": self.bbox,
            "column_count": self.column_count,
            "row_count": self.row_count,
            "headers": self.headers,
            "rows": self.rows,
            "merged_cells": self.merged_cells,
            "metadata": self.metadata,
        }

    def to_markdown(self) -> str:
        """转换为 Markdown 表格格式"""
        if not self.rows:
            return ""
        
        lines = []
        # 标题
        if self.title:
            lines.append(f"**{self.title}**")
        
        # 表头
        headers = self.headers or (self.rows[0] if self.rows else [])
        if headers:
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 数据行
        start_idx = 1 if self.headers else 0
        for row in self.rows[start_idx:]:
            # 处理合并单元格：用空字符串占位
            expanded_row = []
            col = 0
            while col < len(row):
                expanded_row.append(row[col] if row[col] else "")
                # 跳过合并单元格占位
                for r_idx, c_idx, r_span, c_span in self.merged_cells:
                    if r_idx == start_idx - 1 + len(lines) - 2 and c_idx == col:
                        # 这是一个合并单元格的起始位置
                        for _ in range(c_span - 1):
                            expanded_row.append("")
                        break
                col += 1
            lines.append("| " + " | ".join(expanded_row) + " |")
        
        return "\n".join(lines)

    def to_text(self, format: str = "markdown") -> str:
        """
        转换为文本格式
        
        Args:
            format: 'markdown' | 'csv' | 'tsv' | 'structured'
        """
        if format == "csv":
            return self._to_csv()
        elif format == "tsv":
            return self._to_tsv()
        elif format == "structured":
            return self._to_structured_text()
        else:
            return self.to_markdown()

    def _to_csv(self) -> str:
        import csv
        output = io.StringIO()
        writer = csv.writer(output)
        if self.headers:
            writer.writerow(self.headers)
        writer.writerows(self.rows)
        return output.getvalue().strip()

    def _to_tsv(self) -> str:
        lines = []
        if self.headers:
            lines.append("\t".join(self.headers))
        for row in self.rows:
            lines.append("\t".join(row))
        return "\n".join(lines)

    def _to_structured_text(self) -> str:
        """
        结构化文本格式，保留表格层级关系，适合 LLM 理解
        """
        lines = []
        lines.append(f"【表格: {self.title or '无标题'} | 页码: {self.page_number} | {self.column_count}列×{self.row_count}行】")
        
        headers = self.headers or (self.rows[0] if self.rows else [])
        data_start = 1 if self.headers else 0
        
        for i, row in enumerate(self.rows):
            row_label = f"行{i+1}" if i < data_start else f"数据行{i+1-data_start}"
            cells_str = []
            for j, (header, cell) in enumerate(zip(headers, row)):
                if cell.strip():
                    cells_str.append(f"{header}: {cell}")
                else:
                    cells_str.append(f"{header}: (空)")
            lines.append(f"  {row_label}: {' | '.join(cells_str)}")
        
        return "\n".join(lines)


# ── 核心解析器 ────────────────────────────────────────────────────────────────

class TableParser:
    """
    基于 Unstructured 的表格解析器
    
    支持格式：PDF / DOCX / PPTX / PNG / JPG
    功能：表格边界识别、合并单元格处理、跨页表格拼接
    """
    
    def __init__(self):
        self._unstructured_available = False
        self._try_import_unstructured()
    
    def _try_import_unstructured(self) -> None:
        """尝试导入 unstructured，失败则降级到基础方案"""
        try:
            from unstructured.partition.pdf import partition_pdf
            from unstructured.partition.docx import partition_docx
            from unstructured.partition.pptx import partition_pptx
            from unstructured.partition.image import partition_image
            from unstructured.documents.elements import Table, Title
            self._unstructured_available = True
            self._partition_pdf = partition_pdf
            self._partition_docx = partition_docx
            self._partition_pptx = partition_pptx
            self._partition_image = partition_image
            self._ElementTable = Table
            self._ElementTitle = Title
            logger.info("[TableParser] Unstructured 库加载成功")
        except ImportError as e:
            logger.warning(f"[TableParser] Unstructured 库不可用，将使用基础解析方案: {e}")
            self._unstructured_available = False
    
    def parse_file(
        self,
        file_content: bytes,
        file_name: str,
        extract_tables: bool = True,
        include_element_text: bool = True,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """
        解析文件，提取表格和文本块
        
        Args:
            file_content: 文件二进制内容
            file_name: 文件名
            extract_tables: 是否提取表格
            include_element_text: 是否包含非表格文本元素
        
        Returns:
            (tables, text_chunks): 表格列表和文本块列表
        """
        ext = file_name.lower().rsplit(".", 1)[-1]
        
        if ext == "pdf":
            return self.parse_pdf(file_content, file_name, extract_tables, include_element_text)
        elif ext == "docx":
            return self.parse_docx(file_content, file_name, extract_tables, include_element_text)
        elif ext in ("pptx", "ppt"):
            return self.parse_pptx(file_content, file_name, extract_tables, include_element_text)
        elif ext in ("png", "jpg", "jpeg"):
            return self.parse_image(file_content, file_name, extract_tables)
        else:
            return self._parse_fallback(file_content, file_name, ext)
    
    def parse_pdf(
        self,
        file_content: bytes,
        file_name: str = "unknown.pdf",
        extract_tables: bool = True,
        include_element_text: bool = True,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """解析 PDF，提取表格"""
        if self._unstructured_available:
            return self._parse_pdf_unstructured(file_content, file_name, extract_tables, include_element_text)
        else:
            return self._parse_pdf_fallback(file_content, file_name)
    
    def parse_docx(
        self,
        file_content: bytes,
        file_name: str = "unknown.docx",
        extract_tables: bool = True,
        include_element_text: bool = True,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """解析 DOCX，提取表格"""
        if self._unstructured_available:
            return self._parse_docx_unstructured(file_content, file_name, extract_tables, include_element_text)
        else:
            return self._parse_docx_fallback(file_content, file_name)
    
    def parse_pptx(
        self,
        file_content: bytes,
        file_name: str = "unknown.pptx",
        extract_tables: bool = True,
        include_element_text: bool = True,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """解析 PPTX，提取表格"""
        if self._unstructured_available:
            return self._parse_pptx_unstructured(file_content, file_name, extract_tables, include_element_text)
        else:
            return [], []
    
    def parse_image(
        self,
        file_content: bytes,
        file_name: str = "unknown.png",
        extract_tables: bool = True,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """解析图片（使用 OCR 识别表格）"""
        if self._unstructured_available:
            return self._parse_image_unstructured(file_content, file_name, extract_tables)
        else:
            return [], []
    
    # ── Unstructured 解析 ─────────────────────────────────────────────────────
    
    def _parse_pdf_unstructured(
        self,
        file_content: bytes,
        file_name: str,
        extract_tables: bool,
        include_element_text: bool,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """使用 Unstructured 解析 PDF"""
        try:
            elements = self._partition_pdf(
                file=io.BytesIO(file_content),
                extract_tables_in_pdf=True,
                strategy="auto",
                infer_table_structure=True,
            )
            return self._process_elements(elements, file_name, extract_tables, include_element_text)
        except Exception as e:
            logger.error(f"[TableParser] Unstructured PDF 解析失败，降级到基础方案: {e}")
            return self._parse_pdf_fallback(file_content, file_name)
    
    def _parse_docx_unstructured(
        self,
        file_content: bytes,
        file_name: str,
        extract_tables: bool,
        include_element_text: bool,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """使用 Unstructured 解析 DOCX"""
        try:
            elements = self._partition_docx(
                file=io.BytesIO(file_content),
                include_page_breaks=True,
            )
            return self._process_elements(elements, file_name, extract_tables, include_element_text)
        except Exception as e:
            logger.error(f"[TableParser] Unstructured DOCX 解析失败: {e}")
            return self._parse_docx_fallback(file_content, file_name)
    
    def _parse_pptx_unstructured(
        self,
        file_content: bytes,
        file_name: str,
        extract_tables: bool,
        include_element_text: bool,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """使用 Unstructured 解析 PPTX"""
        try:
            elements = self._partition_pptx(
                file=io.BytesIO(file_content),
                include_page_breaks=True,
            )
            return self._process_elements(elements, file_name, extract_tables, include_element_text)
        except Exception as e:
            logger.error(f"[TableParser] Unstructured PPTX 解析失败: {e}")
            return [], []
    
    def _parse_image_unstructured(
        self,
        file_content: bytes,
        file_name: str,
        extract_tables: bool,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """使用 Unstructured 解析图片（OCR）"""
        try:
            elements = self._partition_image(
                file=io.BytesIO(file_content),
                languages=["chi_sim", "eng"],  # 中英文
                strategy="hi_res",  # 高精度模式
                extract_tables_in_pdf=False,
            )
            return self._process_elements(elements, file_name, extract_tables, True)
        except Exception as e:
            logger.error(f"[TableParser] Unstructured 图片 OCR 解析失败: {e}")
            return [], []
    
    # ── 元素处理 ──────────────────────────────────────────────────────────────
    
    def _process_elements(
        self,
        elements,
        file_name: str,
        extract_tables: bool,
        include_element_text: bool,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """处理 Unstructured 元素列表，提取表格和文本"""
        tables = []
        text_chunks = []
        pending_title = ""
        current_page = 0
        
        for elem in elements:
            # 更新页码
            page_number = getattr(elem, 'metadata', {}).get('page_number', 0)
            if page_number:
                current_page = page_number
            
            elem_type = elem.category if hasattr(elem, 'category') else type(elem).__name__
            
            # 提取标题
            if elem_type in ('Title', 'title'):
                pending_title = str(elem)
                continue
            
            # 提取表格
            if extract_tables and elem_type in ('Table', 'table'):
                table = self._extract_table_from_element(elem, file_name, current_page)
                if table:
                    # 如果有等待的标题，关联到表格
                    if pending_title:
                        table.title = pending_title
                        pending_title = ""
                    tables.append(table)
                continue
            
            # 提取文本块
            if include_element_text and elem_type not in ('Table', 'table'):
                text = str(elem).strip()
                if text:
                    text_chunks.append({
                        "content": text,
                        "page_number": current_page,
                        "element_type": elem_type,
                    })
            
            # 重置标题（如果标题后跟的是非表格内容，标题可能不属于任何表格）
            if elem_type not in ('Table', 'table', 'Title', 'title'):
                pending_title = ""
        
        # 跨页表格拼接
        if len(tables) > 1:
            tables = self._merge_cross_page_tables(tables)
        
        logger.info(f"[TableParser] 解析完成: {len(tables)} 个表格, {len(text_chunks)} 个文本块")
        return tables, text_chunks
    
    def _extract_table_from_element(self, element, file_name: str, page_number: int) -> Optional[TableData]:
        """从 Unstructured Table 元素中提取表格数据"""
        try:
            # Unstructured 的 Table 元素有 metadata.table 属性
            metadata = getattr(element, 'metadata', {})
            table_metadata = metadata.get('table', {})
            
            # 获取表格的行列数据
            rows = table_metadata.get('rows', [])
            if not rows:
                # 尝试从 element.text 解析
                return self._parse_table_from_text(str(element), file_name, page_number)
            
            table = TableData(
                file_name=file_name,
                page_number=page_number,
            )
            
            # 提取 bbox
            bbox = metadata.get('coordinates', {})
            if bbox:
                points = bbox.get('points', [])
                if len(points) >= 4:
                    xs = [p[0] for p in points]
                    ys = [p[1] for p in points]
                    table.bbox = (min(xs), min(ys), max(xs), max(ys))
            
            # 处理行数据
            headers = []
            data_rows = []
            merged_cells = []
            
            for row_idx, row in enumerate(rows):
                row_data = []
                col_idx = 0
                for cell in row:
                    cell_text = cell.get('text', '').strip()
                    is_header = cell.get('is_header', False)
                    row_span = cell.get('row_span', 1)
                    col_span = cell.get('col_span', 1)
                    
                    row_data.append(cell_text)
                    
                    if is_header and row_idx == 0:
                        headers.append(cell_text)
                    
                    if row_span > 1 or col_span > 1:
                        merged_cells.append((row_idx, col_idx, row_span, col_span))
                    
                    col_idx += col_span
                
                data_rows.append(row_data)
            
            # 推断列数（取最大行长度）
            col_count = max(len(r) for r in data_rows) if data_rows else 0
            
            # 统一行长度
            for row in data_rows:
                while len(row) < col_count:
                    row.append("")
            
            table.column_count = col_count
            table.row_count = len(data_rows)
            table.headers = headers if headers else (data_rows[0] if data_rows else [])
            table.rows = data_rows
            table.merged_cells = merged_cells
            
            return table
            
        except Exception as e:
            logger.warning(f"[TableParser] 表格元素提取失败: {e}")
            return self._parse_table_from_text(str(element), file_name, page_number)
    
    def _parse_table_from_text(self, text: str, file_name: str, page_number: int) -> Optional[TableData]:
        """从纯文本中解析表格（降级方案）"""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 2:
            return None
        
        # 检测分隔符
        delimiter = self._detect_delimiter(lines)
        if not delimiter:
            return None
        
        # 解析行列
        rows = []
        for line in lines:
            cells = [c.strip() for c in line.split(delimiter)]
            # 过滤掉纯分隔符的行（如 |---|---|）
            if all(c in ('', '-', '---', '---|') for c in cells):
                continue
            if cells:
                rows.append(cells)
        
        if len(rows) < 2:
            return None
        
        table = TableData(
            file_name=file_name,
            page_number=page_number,
        )
        
        col_count = max(len(r) for r in rows)
        for row in rows:
            while len(row) < col_count:
                row.append("")
        
        table.column_count = col_count
        table.row_count = len(rows)
        table.headers = rows[0]
        table.rows = rows[1:]
        
        return table
    
    def _detect_delimiter(self, lines: List[str]) -> Optional[str]:
        """检测表格分隔符"""
        for delim in ['\t', '|', ',']:
            counts = [line.count(delim) for line in lines]
            if counts and min(counts) > 0 and max(counts) == min(counts):
                return delim
        return None
    
    # ── 跨页表格合并 ──────────────────────────────────────────────────────────
    
    def _merge_cross_page_tables(self, tables: List[TableData]) -> List[TableData]:
        """
        检测并合并跨页表格
        
        合并条件：
        1. 相邻表格的表头相同或高度相似
        2. 前一个表格的最后一行和后一个表格的第一行在语义上连续
        3. 表格之间没有其他内容（在原始文档中相邻）
        """
        if len(tables) < 2:
            return tables
        
        merged = [tables[0]]
        
        for i in range(1, len(tables)):
            prev = merged[-1]
            curr = tables[i]
            
            # 检查是否应该合并
            if self._should_merge_tables(prev, curr):
                merged_table = self._merge_two_tables(prev, curr)
                merged[-1] = merged_table
                logger.info(
                    f"[TableParser] 跨页表格合并: page={prev.page_number} → page={curr.page_number}, "
                    f"合并后 {merged_table.row_count} 行"
                )
            else:
                merged.append(curr)
        
        return merged
    
    def _should_merge_tables(self, t1: TableData, t2: TableData) -> bool:
        """判断两个表格是否应该合并"""
        # 条件1：表头相同
        if t1.headers and t2.headers:
            header_sim = self._list_similarity(t1.headers, t2.headers)
            if header_sim >= 0.8:
                return True
        
        # 条件2：t2 的表头是 t1 表头的子集（可能是跨页后只显示了部分表头）
        if t1.headers and t2.headers:
            if set(t2.headers).issubset(set(t1.headers)) and len(t2.headers) >= 2:
                return True
        
        # 条件3：相邻页码且列数相同，且 t2 的第一行看起来像数据行（不是表头）
        if abs(t1.page_number - t2.page_number) <= 1 and t1.column_count == t2.column_count:
            if t1.rows and t2.rows:
                # 如果 t2 有 headers 属性，说明第一行是表头 → 可以合并
                if t2.headers:
                    return True
                # 如果 t2 没有 headers 属性，检查第一行是否是数据行
                first_row = t2.rows[0]
                if not self._looks_like_header(first_row):
                    return True
                # 如果第一行看起来像表头，但和 t1 的表头高度相似，也合并
                if t1.headers:
                    first_row_sim = self._list_similarity(first_row, t1.headers)
                    if first_row_sim >= 0.5:
                        return True
        
        return False
    
    def _merge_two_tables(self, t1: TableData, t2: TableData) -> TableData:
        """合并两个表格"""
        merged = TableData(
            title=t1.title or t2.title,
            page_number=t1.page_number,
            bbox=t1.bbox,
            headers=t1.headers if t1.headers else t2.headers,
            column_count=max(t1.column_count, t2.column_count),
            metadata={**t1.metadata, **t2.metadata, "merged_pages": [t1.page_number, t2.page_number]},
        )
        
        # 决定 t2 的第一行是否为表头（需要跳过）
        skip_first = False
        if t2.headers:
            # t2 明确有 headers 属性 → 第一行是表头
            skip_first = True
        elif t2.rows and t1.headers:
            # t2 没有 headers 属性，检查第一行是否与 t1 的表头相似
            first_row_sim = self._list_similarity(t2.rows[0], t1.headers)
            if first_row_sim >= 0.5:
                skip_first = True
        
        start_idx = 1 if skip_first else 0
        
        merged.rows = list(t1.rows) + list(t2.rows[start_idx:])
        merged.row_count = len(merged.rows)
        
        # 合并合并单元格信息（调整 t2 的行索引）
        merged.merged_cells = list(t1.merged_cells)
        row_offset = len(t1.rows)
        for r, c, rs, cs in t2.merged_cells:
            adjusted_r = r + row_offset
            # 如果跳过了 t2 的第一行，需要调整索引
            if skip_first and r >= 1:
                adjusted_r = r - 1 + row_offset
            merged.merged_cells.append((adjusted_r, c, rs, cs))
        
        return merged
    
    def _list_similarity(self, list1: List[str], list2: List[str]) -> float:
        """计算两个列表的相似度（Jaccard）"""
        if not list1 and not list2:
            return 1.0
        if not list1 or not list2:
            return 0.0
        
        set1 = set(s.strip().lower() for s in list1 if s.strip())
        set2 = set(s.strip().lower() for s in list2 if s.strip())
        
        if not set1 or not set2:
            return 0.0
        
        intersection = set1 & set2
        union = set1 | set2
        
        return len(intersection) / len(union)
    
    def _looks_like_header(self, row: List[str]) -> bool:
        """判断一行是否像表头"""
        if not row:
            return False
        
        # 表头通常包含：单位、类型、分类等关键词
        header_keywords = ['单位', '类型', '名称', '编号', '日期', '时间', '项目', '指标',
                          '金额', '数量', '占比', '比例', '说明', '备注', '备注', '备注']
        
        text = ' '.join(row).lower()
        keyword_matches = sum(1 for kw in header_keywords if kw in text)
        
        # 如果包含多个表头关键词，或者是全大写的英文
        if keyword_matches >= 2:
            return True
        
        # 如果所有单元格都很短（表头通常比较简短）
        avg_len = sum(len(c) for c in row) / len(row)
        if avg_len < 10 and len(row) >= 3:
            return True
        
        return False
    
    # ── 基础降级方案 ──────────────────────────────────────────────────────────
    
    def _parse_pdf_fallback(
        self,
        file_content: bytes,
        file_name: str,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """基础 PDF 表格解析（使用 PyMuPDF）"""
        try:
            import fitz
            doc = fitz.open(stream=file_content, filetype="pdf")
            tables = []
            text_chunks = []
            
            for page_num, page in enumerate(doc, 1):
                page_tables = page.find_tables()
                if page_tables.tables:
                    for tab in page_tables.tables:
                        table = self._convert_pymupdf_table(tab, file_name, page_num)
                        if table:
                            tables.append(table)
                
                # 提取非表格文本
                page_text = page.get_text("text")
                for block in page.get_text("blocks"):
                    text = block[4].strip()
                    if text:
                        text_chunks.append({
                            "content": text,
                            "page_number": page_num,
                            "element_type": "text",
                        })
            
            doc.close()
            
            if len(tables) > 1:
                tables = self._merge_cross_page_tables(tables)
            
            return tables, text_chunks
            
        except Exception as e:
            logger.error(f"[TableParser] PDF 基础解析失败: {e}")
            return [], []
    
    def _convert_pymupdf_table(self, table, file_name: str, page_number: int) -> Optional[TableData]:
        """将 PyMuPDF Table 转换为 TableData"""
        try:
            data = table.extract()
            if not data:
                return None
            
            # 去除空行
            data = [row for row in data if any(cell and cell.strip() for cell in row)]
            if not data:
                return None
            
            col_count = max(len(row) for row in data)
            for row in data:
                while len(row) < col_count:
                    row.append("")
            
            table_data = TableData(
                file_name=file_name,
                page_number=page_number,
                column_count=col_count,
                row_count=len(data),
                headers=data[0] if data else [],
                rows=data[1:] if data else [],
                bbox=table.bbox if hasattr(table, 'bbox') else None,
            )
            
            # 检测合并单元格（PyMuPDF 不直接提供，通过空单元格推断）
            table_data.merged_cells = self._detect_merged_cells_from_pymupdf(data)
            
            return table_data
            
        except Exception as e:
            logger.warning(f"[TableParser] PyMuPDF 表格转换失败: {e}")
            return None
    
    def _detect_merged_cells_from_pymupdf(self, data: List[List[str]]) -> List[Tuple[int, int, int, int]]:
        """
        从 PyMuPDF 表格数据中推断合并单元格
        规则：连续的空单元格可能表示合并
        """
        merged = []
        if not data:
            return merged
        
        rows = len(data)
        cols = max(len(r) for r in data) if data else 0
        
        # 检测横向合并：同一行中，如果一个单元格很长而后面跟着空单元格
        for r in range(rows):
            row = data[r]
            c = 0
            while c < len(row) - 1:
                if row[c] and not row[c + 1].strip():
                    # 检查是否是横向合并
                    span = 1
                    for nc in range(c + 1, len(row)):
                        if not row[nc].strip():
                            span += 1
                        else:
                            break
                    if span > 1:
                        merged.append((r, c, 1, span))
                    c += span
                else:
                    c += 1
        
        return merged
    
    def _parse_docx_fallback(
        self,
        file_content: bytes,
        file_name: str,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """基础 DOCX 表格解析（使用 python-docx）"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))
            
            tables = []
            text_chunks = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_chunks.append({
                        "content": para.text.strip(),
                        "page_number": 0,
                        "element_type": "paragraph",
                    })
            
            for tbl_idx, table in enumerate(doc.tables):
                table_data = self._convert_docx_table(table, file_name, tbl_idx)
                if table_data:
                    tables.append(table_data)
            
            return tables, text_chunks
            
        except Exception as e:
            logger.error(f"[TableParser] DOCX 基础解析失败: {e}")
            return [], []
    
    def _convert_docx_table(self, table, file_name: str, table_index: int) -> Optional[TableData]:
        """将 python-docx Table 转换为 TableData"""
        try:
            rows = []
            merged_cells = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                col_idx = 0
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    
                    # 检测合并单元格
                    tc = cell._tc
                    grid_span = tc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                    if grid_span:
                        span = int(grid_span)
                        if span > 1:
                            merged_cells.append((row_idx, col_idx, 1, span))
                    
                    col_idx += 1
                
                rows.append(row_data)
            
            if not rows:
                return None
            
            col_count = max(len(r) for r in rows)
            for row in rows:
                while len(row) < col_count:
                    row.append("")
            
            return TableData(
                file_name=file_name,
                page_number=0,
                column_count=col_count,
                row_count=len(rows),
                headers=rows[0] if rows else [],
                rows=rows[1:] if rows else [],
                merged_cells=merged_cells,
            )
            
        except Exception as e:
            logger.warning(f"[TableParser] DOCX 表格转换失败: {e}")
            return None
    
    def _parse_fallback(
        self,
        file_content: bytes,
        file_name: str,
        ext: str,
    ) -> Tuple[List[TableData], List[Dict[str, Any]]]:
        """不支持格式的基础解析"""
        try:
            text = file_content.decode("utf-8", errors="ignore")
            return [], [{"content": text, "page_number": 0, "element_type": "text"}]
        except Exception:
            return [], []


# ── 模块级便捷函数 ────────────────────────────────────────────────────────────

_parser_instance: Optional[TableParser] = None


def get_table_parser() -> TableParser:
    """获取单例 TableParser"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = TableParser()
    return _parser_instance


def parse_tables_from_file(
    file_content: bytes,
    file_name: str,
) -> Tuple[List[TableData], List[Dict[str, Any]]]:
    """
    便捷函数：从文件内容中解析表格
    
    Args:
        file_content: 文件二进制内容
        file_name: 文件名
    
    Returns:
        (tables, text_chunks): 表格列表和文本块列表
    """
    parser = get_table_parser()
    return parser.parse_file(file_content, file_name)
